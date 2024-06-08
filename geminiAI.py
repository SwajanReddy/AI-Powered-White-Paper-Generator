import google.generativeai as genai
import time
from docx import Document

def check_net(r):
    if len(r.text)!=0:
        print("response received")
    else:
        print("failed connection to gemini")

def init_bot(t):
    genai.configure(api_key="")     #Insert your own Gemini API Key here
    d1 = Document()
    model=genai.GenerativeModel(
    model_name="gemini-1.5-pro-latest",
    system_instruction=['''You are a research paper writer. You can generate white paper summariers of research papers and conference papers.''',
    '''the typical format of any white paper includes:
    Title:
    You will want a title that will grab attention and communicates the problem you are solving, which is appropriate for your target audience. Including the title’s product name is not a good idea as your title should be benefit-oriented, not feature, oriented.
    The executive summary:
    Include direct, brief statements of your position to engage the reader. You need to provide enough facts to assure and encourage the reader to continue reading the paper even though it summarizes what the document entails.
    Introduction:
    It is an overview of the white paper. It discusses the main points that comprise the document. It also helps the customer identify the relevance of the content to them. In this section, define the issue and provide a background discussion.
    Problem statement:
    The problem statement is the problem faced by customers that the white paper will address. It needs to be stated clearly for easier understanding by the reader. This part should entirely be from the point of view of the target audience.
    High-level solution:
    Including any competing technologies, describe the relevant technologies at a high level. Support the arguments with tables, charts, and graphics. You are educating your target audience on where your solution fits and the current state of the art.
    Solution:
    The solution section helps generate leads for the company as it talks about how the product or service will help the reader. In white papers, the gentle approach works better. Use appropriate language and level of detail according to your target audience. Remember to show why your solution is much better than that of the competition.
    Business benefits:
    In this section, you need to grab the reader’s attention by providing plenty of assurances that your solution will work for them. Show the reader that you understand his/her pain.
    Conclusion:
    The conclusion is a summary of the entire article, and it is advisable to include the essential points. Many readers may skip the paper’s entire contents to read this section only, so it is advisable to write this section as if it were a standalone document summarizing the main selling points about your solution.
    Don't give conclusions in any other section, we have a separate conclusion section in the end for the whole document.
    '''])

    chat = model.start_chat(history=[])
    
    msg='''I am a research professional. I need you to read research paper content that i provide you and generate a white paper
    summary in the following format:
    Title
    You will want a title that will grab attention and communicates the problem you are solving, which is appropriate for your target audience. Including the title’s product name is not a good idea as your title should be benefit-oriented, not feature, oriented.
    The executive summary
    Include direct, brief statements of your position to engage the reader. You need to provide enough facts to assure and encourage the reader to continue reading the paper even though it summarizes what the document entails.
    Introduction
    It is an overview of the white paper. It discusses the main points that comprise the document. It also helps the customer identify the relevance of the content to them. In this section, define the issue and provide a background discussion.
    Problem statement
    The problem statement is the problem faced by customers that the white paper will address. It needs to be stated clearly for easier understanding by the reader. This part should entirely be from the point of view of the target audience.
    High-level solution
    Including any competing technologies, describe the relevant technologies at a high level. Support the arguments with tables, charts, and graphics. You are educating your target audience on where your solution fits and the current state of the art.
    Solution
    The solution section helps generate leads for the company as it talks about how the product or service will help the reader. In white papers, the gentle approach works better. Use appropriate language and level of detail according to your target audience. Remember to show why your solution is much better than that of the competition.
    Business benefits
    In this section, you need to grab the reader’s attention by providing plenty of assurances that your solution will work for them. Show the reader that you understand his/her pain.
    Conclusion
    The conclusion is a summary of the entire article, and it is advisable to include the essential points. Many readers may skip the paper’s entire contents to read this section only, so it is advisable to write this section as if it were a standalone document summarizing the main selling points about your solution.
    Do you understand?
    '''
    r = chat.send_message(msg)
    check_net(r)
    d1.add_paragraph(r.text)
    t.write("Processing... estimated time to complete: 6 min")
    
    msg = '''Alright, first i will give you entire paper content in the format of heading:content.
    then i am gonna ask you to write each of the sections. While giving your responses don't include any headers, just give me the content
    for what i asked for. Am i clear?'''
    r = chat.send_message(msg)
    check_net(r)
    d1.add_paragraph(r.text)
    print("in timeout 60s")
    time.sleep(60)
    print("timeout completed 60s")
    t.write("Processing... estimated time to complete: 5 min")
    d1.save("r1.docx")
    return chat


def p1(headings_and_contents, chat, t, s, b):
    d1 = Document()
    msg = f'''Alright so the paper content is:\n{headings_and_contents}\n
    process this content, I am gonna ask you to generate section wise content as i said earlier in the following messages.
    '''
    r = chat.send_message(msg)
    check_net(r)
    d1.add_paragraph(r.text)
    d1.save("r2.docx")
    t.write("Processing... estimated time to complete: 4.5 min")
    print("in timeout 30s at p1")
    time.sleep(30)
    print("time out completed 30s")
 
    prompts = {
    'Summary': 'Generate a summary for that content that i put in the earlier message:',
    'Introduction': 'Generate an introduction',
    'Problem Statement': 'Generate a problem statement',
    'High-level Solution': 'Generate a high-level solution',
    'Solution': 'Generate a solution for the following text',
    'Business Solutions': 'Identify and describe potential business solutions based on the research paper content',
    'Conclusions': 'Generate conclusions',
    'References': 'Identify and list references'
    }
    
    output_doc = Document()
    msg = "Get me Title"
    r = chat.send_message(msg)
    t.write("Processing... estimated time to complete: 4.5 min")
    s.write("Working on Title")
    x = len(prompts)+1
    print("x:",x)
    b.progress(1 / x)
    print("in time out 30s in p2")
    time.sleep(30)
    print("time out completed 30s")
    tm =  4
    output_doc.add_heading(r.text, level=1)
    for i, (section_name, prompt_text) in enumerate(prompts.items(), start=2):
        t.write(f"Processing... estimated time to complete: {tm} min")
        s.write(f"Working on {section_name}")
        prompt=prompt_text + '\n'
        response = chat.send_message(prompt)
        output_doc.add_heading(section_name, level=2)
        output_doc.add_paragraph(response.text)
        print(f'Finished for section{section_name}')
        print(f"i={i}, x={x}, i/x = {i/x}")
        b.progress(i / x)
        tm = tm - 0.5
        time.sleep(30)
        print("time out completed 30s")
    output_doc.save("ans.docx")
    return output_doc



def run_bot(headings_and_contents, t, s, b):
    chat = init_bot(t)
    print("bot initialized and tone is set.")
    doc = p1(headings_and_contents, chat, t, s, b)
    #doc = p2(chat, t, s, b)
    print("completed run function.")
    print(type(doc))
    return doc, True











